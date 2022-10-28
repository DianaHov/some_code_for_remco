from django.core.management import BaseCommand
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

from cases.models import Management_command_log
from cases.models_client import Upl_file

import io
import fitz
import docx
import pytz
from datetime import datetime

from openpyxl import load_workbook
from openpyxl.worksheet._read_only import ReadOnlyWorksheet
from striprtf.striprtf import rtf_to_text
from email import message_from_file

class Command(BaseCommand):
    """
    This job is for extraction of the content of newly uploaded files
    for further vectorization and full text search.
    Supported types are .pdf, .docx, .xlsx, .xlsm, .rtf, .txt, .csv, .eml.
    The class inherits from BaseCommand so it can be run as a django MC
    (python3 manage.py extractTexts) or used as a cron job.
    There is also a button designed for manual run for demos/tests
    (see Upl_files admin page)
    """
    help = 'running this management command will extract texts from new files.'

    def __init__(self):
        """defining all the mime types here"""
        self.command_name = 'extractTexts'
        # .pdf
        self.pdf_mime_types = ['application/pdf',
                               ]
        # .docx
        self.doc_mime_types = ['application/' 
                               'vnd.openxmlformats-officedocument.' 
                               'wordprocessingml.document',
                               ]
        # .xlsx, .xlsm
        self.excel_mime_types = ['application/' 
                                 'vnd.openxmlformats-officedocument.' 
                                 'spreadsheetml.sheet',
                                 'application/vnd.ms-excel.' 
                                 'sheet.macroenabled.12',
                                ]
        # .rtf
        self.rtf_mime_types = ['application/rtf',
                               'text/rtf',
                               'application/msword' #I know this one looks weird
                               ]
        # .txt
        self.txt_mime_types = ['text/plain',
                               ]
        # .csv
        self.csv_mime_types = ['application/vnd.ms-excel',
                               'application/csv', #I don't like it either
                               'text/csv', #according to standarts
                               ]
        # .eml (emails)
        self.eml_mime_types = ['message/rfc822',
                               ]

        super(Command, self).__init__()

    def _create_new_log(self, options):
        """
        creating a new log so we can track all the runs
        and run statuses through admin console
        (see Management_command_log page)
        """
        command_line = self.command_name
        utc_now = datetime.now(pytz.timezone('UTC'))
        self.log = Management_command_log\
            .objects\
            .create(command_name=self.command_name,
                    command_line=command_line,
                    exception_msg='',
                    start_time=utc_now,
                    )
    def _write_log(self, except_msg):
        """writing exception Trace backs to the log if any"""
        self.log.exception_msg += f'{except_msg} \n'

    def _close_log(self):
        """closing log"""
        self.log.end_time = datetime.now(pytz.timezone('UTC'))
        self.log.succeed = True \
            if not self.log.exception_msg \
            else False
        self.log.save()

    def extract_pdf_files(self, files):
        """pdf extractor"""
        for file in files:
            print(file.fl.name)
            try:
                if not settings.USE_S3:
                    fd = open(file.fl.path, 'rb')
                else:
                    fl = default_storage.open(file.fl.name, 'rb')
                    object_as_streaming_body = fl.obj.get()['Body']
                    fd = object_as_streaming_body.read()
                file.content_extracted = True
                file.save()
                doc = fitz.open("pdf", fd)
                file.content = ''
                for page in doc:
                    file.content += page.get_text()
                file.save()
            except Exception as e:
                print(str(e))
                msg = f'{type(e).__name__} {e}'
                self._write_log(msg)
            finally:
                try:
                    fd.close()
                    fl.close()
                except:
                    pass

    def extract_docx_files(self, files):
        """microsoft docs extractor"""
        for file in files:
            print(file.fl.name)
            try:
                file.content_extracted = True
                file.save()
                if not settings.USE_S3:
                    text = docx.Document(file.fl.path)
                else:
                    fl = default_storage.open(file.fl.name, 'rb')
                    object_as_streaming_body = fl.obj.get()['Body']
                    object_as_bytes = object_as_streaming_body.read()
                    object_as_file_like = io.BytesIO(object_as_bytes)
                    text = docx.Document(object_as_file_like)
                file.content = ''
                for paragraph in text.paragraphs:
                    file.content += f' {paragraph.text}'
                file.save()
            except Exception as e:
                print(str(e))
                self._write_log(f'{type(e).__name__} {e}')
            finally:
                try:
                    fl.close()
                except:
                    pass

    def extract_rtf_files(self, files):
        """rtf files extractor"""
        for file in files:
            print(file.fl.name)
            try:
                file.content_extracted = True
                file.save()
                if not settings.USE_S3:
                    fl = open(file.fl.path, 'rb')
                else:
                    fl = default_storage.open(file.fl.name, 'rb')
                file.content = rtf_to_text(fl.read().decode(errors='replace'))
                file.save()
            except Exception as e:
                print(str(e))
                self._write_log(f'{type(e).__name__} {e}')
            finally:
                try:
                    fl.close()
                except:
                    pass

    def extract_txt_files(self, files):
        """text files extractor"""
        for file in files:
            print(file.fl.name)
            try:
                file.content_extracted = True
                file.save()
                if not settings.USE_S3:
                    fl = open(file.fl.path, 'rb')
                else:
                    fl = default_storage.open(file.fl.name, 'rb')
                file.content = fl.read().decode(errors='replace')
                file.save()
            except Exception as e:
                print(str(e))
                self._write_log(f'{type(e).__name__} {e}')
            finally:
                try:
                    fl.close()
                except:
                    pass

    def extract_csv_files(self, files):
        """csv files extractor"""
        self.extract_txt_files(files)

    def extract_excel_files(self, files):
        """exel files extractor"""
        for file in files:
            print(file.fl.name)
            try:
                file.content_extracted = True
                file.save()
                if not settings.USE_S3:
                    wb = load_workbook(filename=file.fl.path,
                                       read_only=True,
                                       data_only=True,
                                       )
                else:
                    fl = default_storage.open(file.fl.name, 'rb')
                    object_as_streaming_body = fl.obj.get()['Body']
                    object_as_bytes = object_as_streaming_body.read()
                    wb = load_workbook(ContentFile(object_as_bytes),
                                       read_only=True,
                                       data_only=True,
                                       )
                file.content = ''
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    if not isinstance(sheet, ReadOnlyWorksheet):
                        continue
                    row_count = sheet.max_row
                    column_count = sheet.max_column
                    grid = ((x, y) for x in range(1, row_count + 1) for y in range(1, column_count + 1))
                    for x, y in grid:
                        file.content += f'{sheet.cell(row=x, column=y).value}, ' \
                            if sheet.cell(row=x, column=y).value \
                            else ''
                file.save()
            except Exception as e:
                print(str(e))
                self._write_log(f'{type(e).__name__} {e}')
            finally:
                try:
                    fl.close()
                except:
                    pass

    def extract_eml_files(self, files):
        """eml content extractor, recursion is used to get all the text or html contents"""
        def _get_content(msg, content = ''):
            if msg.get_content_type() in ['text/plain',
                                          'text/html',
                                          ]:
                return content + f' {msg.get_payload()}'
            if msg.get_content_type() in['multipart/alternative',
                                         'multipart/related',
                                         'multipart/mixed',
                                         'multipart/digest',
                                         ]:
                for inner in msg.get_payload():
                    return _get_content(inner, content)
            return content

        for file in files:
            try:
                file.content_extracted = True
                file.save()
                if not settings.USE_S3:
                    msg = message_from_file(open(file.fl.path, 'r'))
                else:
                    msg = message_from_file(default_storage.open(file.fl.name, 'r'))
                file.content = _get_content(msg)
                file.save()
            except Exception as e:
                print(str(e))
                self._write_log(f'{type(e).__name__} {e}')

    def handle(self, *args, **options):
        """the method called for MC"""
        self._create_new_log(options)
        upl_files = Upl_file.objects.filter(content_extracted = False,
                                            )
        pdf_files = upl_files.filter(content_type__in=self.pdf_mime_types,
                                     )
        doc_files = upl_files.filter(content_type__in=self.doc_mime_types,
                                     )
        excel_files = upl_files.filter(content_type__in=self.excel_mime_types,
                                     )
        eml_files = upl_files.filter(content_type__in=self.eml_mime_types,
                                     )
        rtf_files = upl_files.filter(content_type__in=self.rtf_mime_types,
                                     fl__endswith = '.rtf'
                                     )
        txt_files = upl_files.filter(content_type__in=self.txt_mime_types,
                                     )
        csv_files = upl_files.filter(content_type__in=self.csv_mime_types,
                                     fl__endswith='.csv',
                                     )
        self.extract_pdf_files(pdf_files)
        self.extract_docx_files(doc_files)
        self.extract_rtf_files(rtf_files)
        self.extract_txt_files(txt_files)
        self.extract_csv_files(csv_files)
        self.extract_excel_files(excel_files)
        self.extract_eml_files(eml_files)
        self._close_log()
        return (f'text extraction finished'
                f'pdf files: {pdf_files.count()}\n'
                f'doc files: {doc_files.count()}\n'
                f'rtf files: {rtf_files.count()}\n'
                f'txt files: {txt_files.count()}\n'
                f'csv files: {csv_files.count()}\n'
                f'excel files: {excel_files.count()}\n'
                f'eml files: {eml_files.count()}\n'
                f'exceptions: {self.log.exception_msg if self.log.exception_msg else 0}'
                )
